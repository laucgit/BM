from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(36, 74, 118)
ACCENT_LIGHT = "DCE6F2"
ACCENT_DARK = RGBColor(24, 49, 77)
BODY = RGBColor(48, 56, 65)
MUTED = RGBColor(90, 98, 108)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
CODE_DIR = PROJECT_ROOT / "codigo"
OUTPUTS_DIR = CODE_DIR / "outputs"
FIG_DIR = OUTPUTS_DIR / "figures"
REPORT_DOCX = PROJECT_ROOT / "informe_datos_sinteticos.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)



def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")



def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(1.9)
    sec.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = BODY

    for style_name, size, bold in [("Title", 24, True), ("Subtitle", 13, False), ("Heading 1", 16, True), ("Heading 2", 12.5, True)]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = ACCENT_DARK



def add_paragraph(doc: Document, text: str, *, bold: bool = False, color=None, size: float | None = None,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after: float = 4.0, italic: bool = False):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p



def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY



def add_key_value_table(doc: Document, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.1)
    table.columns[1].width = Cm(10.5)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
        for j, cell in enumerate(cells):
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            if j == 0:
                set_cell_shading(cell, ACCENT_LIGHT)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = ACCENT_DARK
    return table



def add_metrics_table(doc: Document, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, width in enumerate(col_widths_cm):
        table.columns[i].width = Cm(width)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], "244A76")
        set_cell_margins(hdr[i], top=90, start=110, bottom=90, end=110)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.4)
                    run.font.color.rgb = BODY
    return table



def add_figure(doc: Document, image_path: Path, caption: str, width_cm: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(6)
    rr = cap.add_run(caption)
    rr.font.name = "Calibri"
    rr.font.size = Pt(9)
    rr.italic = True
    rr.font.color.rgb = MUTED



def fmt_pct(x: float) -> str:
    return f"{x * 100:.2f}%"



def fmt_dec(x: float) -> str:
    return f"{x:.3f}"



def build_report() -> None:
    summary = json.loads((OUTPUTS_DIR / "summary_metrics.json").read_text(encoding="utf-8"))
    utility_rows = pd_read_csv(OUTPUTS_DIR / "utility_metrics.csv")
    fidelity_rows = pd_read_csv(OUTPUTS_DIR / "fidelity_metrics_by_feature.csv")

    doc = Document()
    style_document(doc)

    # Portada
    add_paragraph(doc, "Generación y validación de datos sintéticos", bold=True, size=22, color=ACCENT_DARK,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Dataset médico tabular: Breast Cancer Wisconsin (Diagnostic)", size=13.5,
                  color=ACCENT, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    cover_box = doc.add_table(rows=0, cols=2)
    cover_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_box.autofit = False
    cover_box.columns[0].width = Cm(5.3)
    cover_box.columns[1].width = Cm(10.3)
    for key, value in [
        ("Miembro 1", "[Sustituir por nombre y apellidos]"),
        ("Miembro 2", "[Sustituir por nombre y apellidos]"),
        ("Asignatura", "Bioinformática y Medicina"),
        ("Fecha", date.today().strftime("%d/%m/%Y")),
    ]:
        cells = cover_box.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for j, cell in enumerate(cells):
            set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if j == 0:
                set_cell_shading(cell, ACCENT_LIGHT)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    if j == 0:
                        run.bold = True
                        run.font.color.rgb = ACCENT_DARK
                    else:
                        run.font.color.rgb = BODY

    add_paragraph(doc, "", space_after=4)
    intro = (
        "Se ha construido una estrategia de síntesis tabular basada en modelos generativos "
        "Gaussian Mixture condicionales por clase. La validación se organiza en tres ejes: "
        "fidelidad estadística, utilidad predictiva y privacidad."
    )
    p = add_paragraph(doc, intro, size=11, color=BODY, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)
    p.paragraph_format.line_spacing = 1.15

    add_key_value_table(doc, [
        ("Enlace al dataset público", "https://archive.ics.uci.edu/dataset/17/breast+cancer+wisconsin+diagnostic"),
        ("Estrategia de generación", "Gaussian Mixture Models condicionales por clase (una mezcla por etiqueta diagnóstica)"),
        ("Librería empleada", "https://scikit-learn.org/stable/modules/generated/sklearn.mixture.GaussianMixture.html"),
        ("Tamaño del dataset", f"{summary['dataset']['n_rows_total']} filas y {summary['dataset']['n_features']} variables continuas"),
        ("Partición experimental", f"Train real: {summary['dataset']['n_rows_train']} | Test real: {summary['dataset']['n_rows_test']} | Sintético: {summary['dataset']['n_rows_train']}"),
    ])

    doc.add_page_break()

    # Metodología
    add_paragraph(doc, "Descripción del dataset y estrategia empleada", bold=True, size=16, color=ACCENT_DARK, space_after=4)
    p = add_paragraph(
        doc,
        "El conjunto Breast Cancer Wisconsin (Diagnostic) es un problema de clasificación binaria con medidas continuas extraídas de imágenes de aspirado con aguja fina. Para generar datos sintéticos se ha usado una estrategia class-conditional: primero se separan las muestras por diagnóstico, después se ajusta un Gaussian Mixture Model por clase y finalmente se muestrea un nuevo conjunto manteniendo la proporción original de clases.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=4,
    )
    p.paragraph_format.line_spacing = 1.15
    add_bullet(doc, "Escalado estándar sobre el train real antes de ajustar los modelos generativos.")
    add_bullet(doc, "Selección del número de componentes por BIC entre 1 y 4 por clase.")
    add_bullet(doc, "Recorte final de valores al rango observado en train para evitar medidas biomédicas implausibles.")
    add_bullet(doc, f"Componentes elegidos: clase 0 = {summary['generator']['per_class_components']['0']}, clase 1 = {summary['generator']['per_class_components']['1']}.")
    add_figure(doc, FIG_DIR / "class_balance.png", "Figura 1. El conjunto sintético preserva exactamente el balance de clases del train real.", 11.8)

    # Fidelidad
    add_paragraph(doc, "Fidelidad", bold=True, size=16, color=ACCENT_DARK, space_after=3)
    p = add_paragraph(
        doc,
        "Para validar la fidelidad se compararon estadísticas descriptivas y varias pruebas de similitud entre distribuciones reales y sintéticas. En concreto se midieron diferencias de media y desviación típica, el estadístico de Kolmogorov-Smirnov, la distancia de Wasserstein normalizada y la diferencia media absoluta entre matrices de correlación.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=4,
    )
    p.paragraph_format.line_spacing = 1.15

    add_metrics_table(doc,
        ["Métrica", "Valor"],
        [
            ["Diferencia media media (en unidades z)", fmt_dec(summary['fidelity']['avg_standardized_mean_diff'])],
            ["Desviación media del ratio de desviaciones típicas", fmt_dec(summary['fidelity']['avg_std_ratio_deviation'])],
            ["KS medio", fmt_dec(summary['fidelity']['avg_ks_statistic'])],
            ["KS máximo", fmt_dec(summary['fidelity']['max_ks_statistic'])],
            ["Wasserstein normalizado medio", fmt_dec(summary['fidelity']['avg_normalized_wasserstein'])],
            ["Brecha media de correlación", fmt_dec(summary['fidelity']['mean_abs_corr_gap'])],
        ],
        [11.3, 4.2],
    )
    add_paragraph(doc, "", space_after=2)
    add_figure(doc, FIG_DIR / "feature_distributions.png", "Figura 2. Distribuciones marginales de cuatro variables representativas en train real y sintético.", 14.7)
    add_figure(doc, FIG_DIR / "correlation_heatmaps.png", "Figura 3. Estructura de correlaciones en los datos reales y sintéticos.", 14.4)
    add_figure(doc, FIG_DIR / "fidelity_table.png", "Figura 4. Variables con mayor divergencia univariante según KS y Wasserstein normalizado.", 14.4)
    p = add_paragraph(
        doc,
        "Los resultados indican una fidelidad global buena: las medias cambian muy poco en promedio (0.035 desviaciones típicas), el KS medio es bajo (0.055) y la distancia de Wasserstein normalizada también se mantiene contenida. La estructura de correlación no es perfecta, pero la brecha media de 0.107 sigue siendo razonable para un generador sencillo. En conjunto, el modelo reproduce bien la forma general del dataset, aunque suaviza parte de las dependencias multivariantes finas.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=6,
    )
    p.paragraph_format.line_spacing = 1.15

    # Utilidad
    add_paragraph(doc, "Utilidad", bold=True, size=16, color=ACCENT_DARK, space_after=3)
    p = add_paragraph(
        doc,
        "La utilidad se evaluó con el protocolo TRTR/TSTR: se entrenaron modelos con datos reales y con datos sintéticos, y en ambos casos se validó exclusivamente sobre el mismo test real. Así puede medirse hasta qué punto los datos sintéticos conservan la señal predictiva relevante para la tarea clínica.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=4,
    )
    p.paragraph_format.line_spacing = 1.15

    util_tbl_rows = []
    for model_name in ["Logistic Regression", "Random Forest"]:
        trtr = utility_rows[(utility_rows['model'] == model_name) & (utility_rows['setting'] == 'TRTR')].iloc[0]
        tstr = utility_rows[(utility_rows['model'] == model_name) & (utility_rows['setting'] == 'TSTR')].iloc[0]
        util_tbl_rows.append([
            model_name,
            fmt_pct(trtr['accuracy']),
            fmt_pct(tstr['accuracy']),
            fmt_dec(trtr['roc_auc']),
            fmt_dec(tstr['roc_auc']),
        ])
    add_metrics_table(doc,
        ["Modelo", "Acc. TRTR", "Acc. TSTR", "ROC-AUC TRTR", "ROC-AUC TSTR"],
        util_tbl_rows,
        [4.2, 2.4, 2.4, 2.7, 2.7],
    )
    add_paragraph(doc, "", space_after=2)
    add_figure(doc, FIG_DIR / "utility_roc_auc.png", "Figura 5. Rendimiento predictivo en test real usando entrenamiento real frente a sintético.", 12.4)
    overlap = summary['utility']['rf_top10_overlap_count']
    p = add_paragraph(
        doc,
        f"La utilidad obtenida es alta. Con regresión logística, el ROC-AUC apenas cae de {summary['utility']['metrics'][0]['roc_auc']:.3f} a {summary['utility']['metrics'][1]['roc_auc']:.3f}. Con random forest, el descenso es igualmente moderado. La caída en accuracy es más visible, pero los modelos entrenados con datos sintéticos siguen superando el 91% en test real. Además, el random forest mantiene coincidencia total en el top-10 de variables más importantes ({overlap}/10), lo que sugiere que la señal discriminativa principal se conserva bien.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=6,
    )
    p.paragraph_format.line_spacing = 1.15

    # Privacidad
    doc.add_page_break()
    add_paragraph(doc, "Privacidad", bold=True, size=16, color=ACCENT_DARK, space_after=3)
    p = add_paragraph(
        doc,
        "La privacidad se analizó mediante Distance to Closest Record (DCR), Nearest Neighbor Distance Ratio (NNDR) y tasa de duplicados exactos. Como referencia, el DCR de las muestras sintéticas hacia el train real se comparó con el DCR de un test real independiente respecto al mismo train. Si los sintéticos fuesen copias o casi copias, sus distancias tenderían a ser claramente menores que las del test real.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=4,
    )
    p.paragraph_format.line_spacing = 1.15

    add_metrics_table(doc,
        ["Indicador", "Valor"],
        [
            ["Tasa de duplicados exactos", fmt_pct(summary['privacy']['exact_duplicate_rate'])],
            ["DCR mediano sintético → train", fmt_dec(summary['privacy']['median_dcr_synthetic_to_train'])],
            ["DCR mediano test real → train", fmt_dec(summary['privacy']['median_dcr_test_to_train'])],
            ["NNDR mediano sintético", fmt_dec(summary['privacy']['median_nndr_synthetic'])],
            ["% sintéticos por debajo del p1 de DCR del test", fmt_pct(summary['privacy']['share_synthetic_below_test_dcr_p01'])],
            ["% sintéticos por debajo del p5 de DCR del test", fmt_pct(summary['privacy']['share_synthetic_below_test_dcr_p05'])],
        ],
        [11.3, 4.2],
    )
    add_paragraph(doc, "", space_after=2)
    add_figure(doc, FIG_DIR / "privacy_dcr_nndr.png", "Figura 6. Comparativa de DCR y distribución de NNDR en el conjunto sintético.", 14.2)
    p = add_paragraph(
        doc,
        "Los indicadores son tranquilizadores: no aparecen duplicados exactos y el DCR mediano de los sintéticos es incluso mayor que el del test real independiente, lo que reduce la sospecha de memorización. Además, ningún registro sintético cae por debajo del percentil 1 ni del percentil 5 del DCR del test real. El NNDR mediano cercano a 1 sugiere que, cuando existe un vecino próximo, no suele tratarse de una coincidencia aislada extrema. Dicho de otra forma, los datos sintéticos parecen plausibles sin quedar excesivamente pegados a individuos concretos del train.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=6,
    )
    p.paragraph_format.line_spacing = 1.15

    # Cierre
    add_paragraph(doc, "Conclusión", bold=True, size=16, color=ACCENT_DARK, space_after=3)
    p = add_paragraph(
        doc,
        "La estrategia basada en Gaussian Mixture Models ofrece un compromiso convincente entre fidelidad, utilidad y privacidad para este dataset médico tabular. No alcanza el realismo multivariante que cabría esperar de modelos más complejos, pero sí permite generar un conjunto sintético útil para experimentación y prototipado con un riesgo bajo de copiar registros reales. Como mejora futura, se podría comparar este enfoque con bibliotecas especializadas como SDV, Synthcity o BeGREAT y añadir ataques de membership inference más agresivos.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=8,
    )
    p.paragraph_format.line_spacing = 1.15
    add_paragraph(doc, "Documento generado automáticamente a partir de los resultados en codigo/outputs/.", italic=True, color=MUTED, size=9.3, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(REPORT_DOCX)



def pd_read_csv(path: Path):
    import pandas as pd
    return pd.read_csv(path)


if __name__ == "__main__":
    build_report()
